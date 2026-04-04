# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2026 Assistance Micro Design
"""Generateur de fichiers Word (.docx) a partir de contenu Markdown."""

from __future__ import annotations

import asyncio
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.core.exceptions import WordGenerationError
from src.models.api import CreateWordParams, CreateWordResult
from src.services.document.base_generator import BaseDocumentGenerator


if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph


logger = logging.getLogger(__name__)


# === Markdown AST ===


@dataclass
class InlineRun:
    """Fragment de texte avec formatage inline."""

    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


@dataclass
class HeadingBlock:
    """Bloc heading (# à ######)."""

    level: int
    text: str


@dataclass
class ParagraphBlock:
    """Bloc paragraphe avec runs inline."""

    runs: list[InlineRun] = field(default_factory=list)


@dataclass
class ListItemBlock:
    """Item de liste (bullet ou numbered)."""

    ordered: bool
    level: int = 0
    runs: list[InlineRun] = field(default_factory=list)


@dataclass
class TableBlock:
    """Bloc table Markdown."""

    headers: list[str]
    rows: list[list[str]]


@dataclass
class CodeBlock:
    """Bloc de code (triple backtick)."""

    content: str
    language: str = ""


@dataclass
class BlockquoteBlock:
    """Bloc citation (> text)."""

    text: str


@dataclass
class PageBreakBlock:
    """Saut de page (---)."""


# Type union pour tous les blocs
Block = (
    HeadingBlock
    | ParagraphBlock
    | ListItemBlock
    | TableBlock
    | CodeBlock
    | BlockquoteBlock
    | PageBreakBlock
)


# === Regex patterns ===

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^(\s*)\d+\.\s+(.+)$")
_HR_RE = re.compile(r"^-{3,}$")
_BLOCKQUOTE_RE = re.compile(r"^>\s*(.*)$")
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_CODE_FENCE_RE = re.compile(r"^```(\w*)$")

# Inline: code > bold italic > bold > italic (ordre de priorite)
_INLINE_RE = re.compile(
    r"(`[^`]+`)"
    r"|(\*{3}.+?\*{3})"
    r"|(\*{2}.+?\*{2})"
    r"|(\*[^*]+?\*)"
)


class WordGenerator(BaseDocumentGenerator):
    """Genere des fichiers Word a partir de contenu Markdown."""

    _error_class = WordGenerationError

    def __init__(self, output_path: Path | None = None) -> None:
        super().__init__(output_path)

    async def generate(self, params: CreateWordParams) -> CreateWordResult:
        """Point d'entree principal. Cree le fichier docx."""
        return await asyncio.to_thread(self._generate_sync, params)

    def _generate_sync(self, params: CreateWordParams) -> CreateWordResult:
        """Generation synchrone du fichier (appelee via to_thread)."""
        self.ensure_output_dir()
        safe_filename = self.sanitize_filename(params.filename)
        file_path = self._output_path / safe_filename
        overwritten = file_path.exists()

        blocks = self._parse_markdown(params.content)
        doc = self._render_document(blocks)

        if params.title:
            doc.core_properties.title = params.title
        if params.author:
            doc.core_properties.author = params.author

        file_size = self._save_and_verify(doc, file_path, safe_filename)

        logger.info(
            "Word genere: %s (%d blocs, %d octets)",
            safe_filename,
            len(blocks),
            file_size,
        )

        return CreateWordResult(
            file_path=str(file_path),
            filename=safe_filename,
            file_size_bytes=file_size,
            overwritten=overwritten,
        )

    # === Markdown Parser ===

    def _parse_markdown(self, content: str) -> list[Block]:
        """Parse le contenu Markdown en liste de blocs types."""
        lines = content.split("\n")
        blocks: list[Block] = []
        i = 0

        while i < len(lines):
            line = lines[i]

            # Blocs multi-lignes (code fence, table)
            fence_match = _CODE_FENCE_RE.match(line)
            if fence_match:
                i, code_blk = self._parse_code_block(lines, i, fence_match.group(1))
                blocks.append(code_blk)
                continue

            if (
                _TABLE_ROW_RE.match(line)
                and i + 1 < len(lines)
                and _TABLE_SEP_RE.match(lines[i + 1])
            ):
                i, tbl_blk = self._parse_table(lines, i)
                blocks.append(tbl_blk)
                continue

            # Blocs single-line
            block = self._parse_line(line)
            if block:
                blocks.append(block)
            i += 1

        return blocks

    def _parse_line(self, line: str) -> Block | None:  # noqa: PLR0911
        """Parse une ligne individuelle en bloc (ou None si ligne vide)."""
        heading_match = _HEADING_RE.match(line)
        if heading_match:
            return HeadingBlock(
                level=len(heading_match.group(1)),
                text=heading_match.group(2).strip(),
            )

        if _HR_RE.match(line.strip()):
            return PageBreakBlock()

        bq_match = _BLOCKQUOTE_RE.match(line)
        if bq_match:
            return BlockquoteBlock(text=bq_match.group(1).strip())

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            return self._parse_list_item(bullet_match, ordered=False)

        num_match = _NUMBERED_RE.match(line)
        if num_match:
            return self._parse_list_item(num_match, ordered=True)

        if not line.strip():
            return None

        return ParagraphBlock(runs=self._parse_inline(line))

    def _parse_list_item(self, match: re.Match[str], *, ordered: bool) -> ListItemBlock:
        """Cree un ListItemBlock a partir d'un match regex."""
        indent = len(match.group(1))
        level = min(indent // 2, 2)
        return ListItemBlock(
            ordered=ordered,
            level=level,
            runs=self._parse_inline(match.group(2)),
        )

    def _parse_code_block(
        self, lines: list[str], start: int, language: str
    ) -> tuple[int, CodeBlock]:
        """Parse un bloc de code delimitee par triple backtick."""
        code_lines: list[str] = []
        i = start + 1
        while i < len(lines):
            if _CODE_FENCE_RE.match(lines[i]):
                i += 1
                break
            code_lines.append(lines[i])
            i += 1
        return i, CodeBlock(content="\n".join(code_lines), language=language)

    def _parse_table(self, lines: list[str], start: int) -> tuple[int, TableBlock]:
        """Parse une table Markdown (header | separator | rows)."""
        header_line = lines[start]
        headers = [c.strip() for c in header_line.strip("|").split("|")]

        i = start + 2  # skip separator
        rows: list[list[str]] = []
        while i < len(lines):
            row_match = _TABLE_ROW_RE.match(lines[i])
            if not row_match:
                break
            cells = [c.strip() for c in lines[i].strip("|").split("|")]
            rows.append(cells)
            i += 1

        return i, TableBlock(headers=headers, rows=rows)

    def _parse_inline(self, text: str) -> list[InlineRun]:
        """Parse les elements inline (bold, italic, code) d'une ligne."""
        runs: list[InlineRun] = []

        last_end = 0
        for match in _INLINE_RE.finditer(text):
            # Texte avant le match
            if match.start() > last_end:
                runs.append(InlineRun(text=text[last_end : match.start()]))

            matched = match.group()
            if match.group(1):  # inline code
                runs.append(InlineRun(text=matched[1:-1], code=True))
            elif match.group(2):  # bold italic
                runs.append(InlineRun(text=matched[3:-3], bold=True, italic=True))
            elif match.group(3):  # bold
                runs.append(InlineRun(text=matched[2:-2], bold=True))
            elif match.group(4):  # italic
                runs.append(InlineRun(text=matched[1:-1], italic=True))

            last_end = match.end()

        # Texte restant
        if last_end < len(text):
            runs.append(InlineRun(text=text[last_end:]))

        if not runs:
            runs.append(InlineRun(text=text))

        return runs

    # === Document Renderer ===

    def _render_document(self, blocks: list[Block]) -> Document:
        """Transforme les blocs parses en document python-docx."""
        doc = Document()

        for block in blocks:
            self._render_block(doc, block)

        return doc

    def _render_block(self, doc: Document, block: Block) -> None:
        """Rend un bloc individuel dans le document."""
        if isinstance(block, HeadingBlock):
            doc.add_heading(block.text, level=block.level)
            return
        if isinstance(block, ParagraphBlock):
            self._add_runs(doc.add_paragraph(), block.runs)
            return
        if isinstance(block, ListItemBlock):
            self._render_list_item(doc, block)
            return
        if isinstance(block, TableBlock):
            self._render_table(doc, block)
            return
        if isinstance(block, CodeBlock):
            self._render_code_block(doc, block)
            return
        if isinstance(block, BlockquoteBlock):
            doc.add_paragraph(block.text, style="Intense Quote")
            return
        if isinstance(block, PageBreakBlock):
            self._render_page_break(doc)

    def _add_runs(self, paragraph: Paragraph, runs: list[InlineRun]) -> None:
        """Ajoute des runs formates a un paragraphe."""
        for inline_run in runs:
            run = paragraph.add_run(inline_run.text)
            if inline_run.bold:
                run.bold = True
            if inline_run.italic:
                run.italic = True
            if inline_run.code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    def _render_list_item(self, doc: Document, block: ListItemBlock) -> None:
        """Rend un item de liste avec le bon style selon le niveau."""
        base_style = "List Number" if block.ordered else "List Bullet"
        suffix = f" {block.level + 1}" if block.level > 0 else ""
        style_name = f"{base_style}{suffix}"

        para = doc.add_paragraph(style=style_name)
        self._add_runs(para, block.runs)

    def _render_table(self, doc: Document, block: TableBlock) -> None:
        """Rend une table Markdown en table Word."""
        num_cols = len(block.headers)
        table = doc.add_table(rows=1, cols=num_cols, style="Table Grid")

        # Headers
        for i, header in enumerate(block.headers):
            table.rows[0].cells[i].text = header

        # Data rows
        for row_data in block.rows:
            row = table.add_row()
            for i, cell in enumerate(row_data):
                if i < num_cols:
                    row.cells[i].text = cell

    def _render_code_block(self, doc: Document, block: CodeBlock) -> None:
        """Rend un bloc de code en paragraphe monospace."""
        para = doc.add_paragraph()
        run = para.add_run(block.content)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

        # Background gris clair via shading
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), "F2F2F2")
        para.paragraph_format.element.get_or_add_pPr().append(shading)

    def _render_page_break(self, doc: Document) -> None:
        """Ajoute un saut de page."""
        para = doc.add_paragraph()
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._element.append(br)

    # === File operations ===

    def _save_and_verify(self, doc: Document, file_path: Path, filename: str) -> int:
        """Sauvegarde le document et verifie la taille."""
        doc.save(str(file_path))
        return self.verify_file_size(file_path, filename)

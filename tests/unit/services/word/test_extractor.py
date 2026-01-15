"""Tests unitaires pour l'extracteur Word."""

from __future__ import annotations

from pathlib import Path

import pytest

from src.models.word import (
    ContentBlock,
    ContentType,
    HeadingLevel,
    WordDocument,
    WordList,
    WordListItem,
    WordParagraph,
    WordTable,
    WordTableCell,
)
from src.services.word.extractor import WordExtractor


class TestWordExtractor:
    """Tests pour l'extracteur Word."""

    @pytest.fixture
    def extractor(self) -> WordExtractor:
        """Crée un extracteur Word."""
        return WordExtractor(extract_images=True)

    @pytest.fixture
    def sample_docx(self, tmp_path: Path) -> Path:
        """Crée un fichier Word de test."""
        from docx import Document  # noqa: PLC0415

        doc = Document()

        # Titre
        doc.add_heading("Document de Test", level=1)

        # Paragraphe
        doc.add_paragraph("Ceci est un paragraphe de test.")

        # Sous-titre
        doc.add_heading("Section 1", level=2)
        doc.add_paragraph("Contenu de la section 1.")

        # Tableau
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Nom"
        table.cell(0, 1).text = "Valeur"
        table.cell(1, 0).text = "Item A"
        table.cell(1, 1).text = "100"
        table.cell(2, 0).text = "Item B"
        table.cell(2, 1).text = "200"

        file_path = tmp_path / "test.docx"
        doc.save(file_path)

        return file_path

    @pytest.mark.asyncio
    async def test_extract_docx_basic(
        self,
        extractor: WordExtractor,
        sample_docx: Path,
    ) -> None:
        """Test extraction basique d'un fichier docx."""
        doc = await extractor.extract(sample_docx)

        assert isinstance(doc, WordDocument)
        assert doc.filename == "test.docx"
        assert len(doc.content_blocks) > 0

    @pytest.mark.asyncio
    async def test_extract_paragraphs(
        self,
        extractor: WordExtractor,
        sample_docx: Path,
    ) -> None:
        """Test extraction des paragraphes."""
        doc = await extractor.extract(sample_docx)

        assert len(doc.paragraphs) > 0

        # Vérifier qu'on a du contenu
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "paragraphe de test" in all_text.lower()

    @pytest.mark.asyncio
    async def test_extract_tables(
        self,
        extractor: WordExtractor,
        sample_docx: Path,
    ) -> None:
        """Test extraction des tableaux."""
        doc = await extractor.extract(sample_docx)

        assert len(doc.tables) >= 1

        # Vérifier le contenu du tableau
        table = doc.tables[0]
        assert table.row_count >= 2

        # Vérifier le Markdown
        md = table.to_markdown()
        assert "Nom" in md
        assert "Valeur" in md

    @pytest.mark.asyncio
    async def test_content_order(
        self,
        extractor: WordExtractor,
        sample_docx: Path,
    ) -> None:
        """Test que l'ordre du contenu est préservé."""
        doc = await extractor.extract(sample_docx)

        orders = [block.order for block in doc.content_blocks]
        assert orders == sorted(orders), "L'ordre doit être croissant"

    @pytest.mark.asyncio
    async def test_to_markdown(
        self,
        extractor: WordExtractor,
        sample_docx: Path,
    ) -> None:
        """Test conversion en Markdown."""
        doc = await extractor.extract(sample_docx)

        md = doc.to_markdown()

        assert "test.docx" in md
        assert len(md) > 100

    @pytest.mark.asyncio
    async def test_file_not_found(self, extractor: WordExtractor) -> None:
        """Test erreur fichier non trouvé."""
        with pytest.raises(FileNotFoundError):
            await extractor.extract("/non/existent/file.docx")

    @pytest.mark.asyncio
    async def test_unsupported_format(
        self,
        extractor: WordExtractor,
        tmp_path: Path,
    ) -> None:
        """Test erreur format non supporté."""
        bad_file = tmp_path / "test.doc"
        bad_file.write_text("not docx")

        with pytest.raises(ValueError, match="Format non supporté"):
            await extractor.extract(bad_file)

    @pytest.mark.asyncio
    async def test_word_count(
        self,
        extractor: WordExtractor,
        sample_docx: Path,
    ) -> None:
        """Test comptage des mots."""
        doc = await extractor.extract(sample_docx)

        assert doc.word_count > 0


class TestWordParagraph:
    """Tests pour WordParagraph."""

    def test_to_markdown_body(self) -> None:
        """Test conversion d'un paragraphe corps."""
        para = WordParagraph(text="Simple texte", level=HeadingLevel.BODY)
        assert para.to_markdown() == "Simple texte"

    def test_to_markdown_heading_1(self) -> None:
        """Test conversion d'un titre niveau 1."""
        para = WordParagraph(text="Titre", level=HeadingLevel.HEADING_1)
        assert para.to_markdown() == "# Titre"

    def test_to_markdown_heading_2(self) -> None:
        """Test conversion d'un titre niveau 2."""
        para = WordParagraph(text="Sous-titre", level=HeadingLevel.HEADING_2)
        assert para.to_markdown() == "## Sous-titre"

    def test_to_markdown_heading_3(self) -> None:
        """Test conversion d'un titre niveau 3."""
        para = WordParagraph(text="Section", level=HeadingLevel.HEADING_3)
        assert para.to_markdown() == "### Section"


class TestWordTable:
    """Tests pour WordTable."""

    def test_to_markdown(self) -> None:
        """Test conversion d'un tableau en Markdown."""
        table = WordTable(
            rows=[
                [WordTableCell(text="A"), WordTableCell(text="B")],
                [WordTableCell(text="1"), WordTableCell(text="2")],
            ]
        )

        md = table.to_markdown()
        assert "| A | B |" in md
        assert "| 1 | 2 |" in md
        assert "| --- | --- |" in md

    def test_to_markdown_empty(self) -> None:
        """Test tableau vide."""
        table = WordTable()
        assert table.to_markdown() == ""

    def test_row_count(self) -> None:
        """Test comptage des lignes."""
        table = WordTable(
            rows=[
                [WordTableCell(text="A")],
                [WordTableCell(text="B")],
                [WordTableCell(text="C")],
            ]
        )
        assert table.row_count == 3

    def test_col_count(self) -> None:
        """Test comptage des colonnes."""
        table = WordTable(
            rows=[
                [WordTableCell(text="A"), WordTableCell(text="B"), WordTableCell(text="C")],
            ]
        )
        assert table.col_count == 3

    def test_to_dict(self) -> None:
        """Test conversion en dictionnaire."""
        table = WordTable(
            rows=[
                [WordTableCell(text="Header1"), WordTableCell(text="Header2")],
                [WordTableCell(text="Val1"), WordTableCell(text="Val2")],
            ]
        )

        d = table.to_dict()
        assert d["row_count"] == 2
        assert d["col_count"] == 2
        assert "Header1" in d["headers"]


class TestWordList:
    """Tests pour WordList."""

    def test_to_markdown_unordered(self) -> None:
        """Test liste à puces."""
        word_list = WordList(
            items=[
                WordListItem(text="Premier"),
                WordListItem(text="Deuxième"),
            ],
            is_ordered=False,
        )

        md = word_list.to_markdown()
        assert "- Premier" in md
        assert "- Deuxième" in md

    def test_to_markdown_ordered(self) -> None:
        """Test liste numérotée."""
        word_list = WordList(
            items=[
                WordListItem(text="Premier"),
                WordListItem(text="Deuxième"),
            ],
            is_ordered=True,
        )

        md = word_list.to_markdown()
        assert "1. Premier" in md
        assert "2. Deuxième" in md

    def test_to_markdown_nested(self) -> None:
        """Test liste avec indentation."""
        word_list = WordList(
            items=[
                WordListItem(text="Parent", level=0),
                WordListItem(text="Enfant", level=1),
            ],
            is_ordered=False,
        )

        md = word_list.to_markdown()
        assert "- Parent" in md
        assert "  - Enfant" in md


class TestContentBlock:
    """Tests pour ContentBlock."""

    def test_paragraph_to_markdown(self) -> None:
        """Test bloc paragraphe."""
        block = ContentBlock(
            content_type=ContentType.PARAGRAPH,
            order=0,
            paragraph=WordParagraph(text="Test"),
        )

        assert block.to_markdown() == "Test"

    def test_table_to_markdown(self) -> None:
        """Test bloc tableau."""
        block = ContentBlock(
            content_type=ContentType.TABLE,
            order=0,
            table=WordTable(
                rows=[
                    [WordTableCell(text="A"), WordTableCell(text="B")],
                ]
            ),
        )

        md = block.to_markdown()
        assert "| A | B |" in md

    def test_empty_block(self) -> None:
        """Test bloc sans contenu."""
        block = ContentBlock(
            content_type=ContentType.PARAGRAPH,
            order=0,
        )

        assert block.to_markdown() == ""


class TestWordDocument:
    """Tests pour WordDocument."""

    def test_to_markdown(self) -> None:
        """Test conversion document complet."""
        doc = WordDocument(
            filename="test.docx",
            file_path="/path/test.docx",
            content_blocks=[
                ContentBlock(
                    content_type=ContentType.HEADING,
                    order=0,
                    paragraph=WordParagraph(text="Titre", level=HeadingLevel.HEADING_1),
                ),
                ContentBlock(
                    content_type=ContentType.PARAGRAPH,
                    order=1,
                    paragraph=WordParagraph(text="Contenu"),
                ),
            ],
            paragraphs=[
                WordParagraph(text="Titre", level=HeadingLevel.HEADING_1),
                WordParagraph(text="Contenu"),
            ],
        )

        md = doc.to_markdown()
        assert "# test.docx" in md
        assert "# Titre" in md
        assert "Contenu" in md

    def test_get_headings(self) -> None:
        """Test récupération des titres."""
        doc = WordDocument(
            filename="test.docx",
            file_path="/path/test.docx",
            paragraphs=[
                WordParagraph(text="Titre 1", level=HeadingLevel.HEADING_1),
                WordParagraph(text="Normal", level=HeadingLevel.BODY),
                WordParagraph(text="Titre 2", level=HeadingLevel.HEADING_2),
            ],
        )

        headings = doc.get_headings()
        assert len(headings) == 2
        assert headings[0].text == "Titre 1"
        assert headings[1].text == "Titre 2"

    def test_get_table_of_contents(self) -> None:
        """Test génération table des matières."""
        doc = WordDocument(
            filename="test.docx",
            file_path="/path/test.docx",
            paragraphs=[
                WordParagraph(text="Introduction", level=HeadingLevel.HEADING_1),
                WordParagraph(text="Section A", level=HeadingLevel.HEADING_2),
            ],
        )

        toc = doc.get_table_of_contents()
        assert "Table des matières" in toc
        assert "Introduction" in toc
        assert "Section A" in toc

# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2025-2026 Assistance Micro Design
"""Tests unitaires pour WordGenerator."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from docx.oxml.ns import qn

from src.core.exceptions import WordGenerationError
from src.models.api import CreateWordParams
from src.services.word.generator import WordGenerator


@pytest.fixture
def generator(tmp_path: Path) -> WordGenerator:
    """WordGenerator avec output_path temporaire."""
    return WordGenerator(output_path=tmp_path)


def _params(
    filename: str = "test.docx",
    content: str = "Hello world",
    title: str | None = None,
    author: str | None = None,
) -> CreateWordParams:
    """Factory pour CreateWordParams."""
    return CreateWordParams(
        filename=filename,
        content=content,
        title=title,
        author=author,
    )


# === Headings ===


@pytest.mark.asyncio
async def test_heading_level_1(generator: WordGenerator) -> None:
    """# Title → Heading level 1."""
    result = await generator.generate(_params(content="# Mon Titre"))

    doc = Document(result.file_path)
    assert doc.paragraphs[0].text == "Mon Titre"
    assert doc.paragraphs[0].style.name == "Heading 1"


@pytest.mark.asyncio
async def test_heading_level_2(generator: WordGenerator) -> None:
    """## Subtitle → Heading level 2."""
    result = await generator.generate(_params(content="## Sous-titre"))

    doc = Document(result.file_path)
    assert doc.paragraphs[0].text == "Sous-titre"
    assert doc.paragraphs[0].style.name == "Heading 2"


@pytest.mark.asyncio
async def test_heading_level_3_to_6(generator: WordGenerator) -> None:
    """### à ###### → Heading 3 à 6."""
    content = "### H3\n#### H4\n##### H5\n###### H6"
    result = await generator.generate(_params(content=content))

    doc = Document(result.file_path)
    for i, level in enumerate(range(3, 7)):
        assert doc.paragraphs[i].style.name == f"Heading {level}"


# === Paragraphs with inline formatting ===


@pytest.mark.asyncio
async def test_bold_text(generator: WordGenerator) -> None:
    """**bold** → run bold."""
    result = await generator.generate(_params(content="Texte **gras** ici"))

    doc = Document(result.file_path)
    runs = doc.paragraphs[0].runs
    bold_runs = [r for r in runs if r.bold]
    assert any(r.text == "gras" for r in bold_runs)


@pytest.mark.asyncio
async def test_italic_text(generator: WordGenerator) -> None:
    """*italic* → run italic."""
    result = await generator.generate(_params(content="Texte *italique* ici"))

    doc = Document(result.file_path)
    runs = doc.paragraphs[0].runs
    italic_runs = [r for r in runs if r.italic]
    assert any(r.text == "italique" for r in italic_runs)


@pytest.mark.asyncio
async def test_bold_italic_text(generator: WordGenerator) -> None:
    """***bold italic*** → run bold + italic."""
    result = await generator.generate(_params(content="Texte ***gras italique*** fin"))

    doc = Document(result.file_path)
    runs = doc.paragraphs[0].runs
    bi_runs = [r for r in runs if r.bold and r.italic]
    assert any("gras italique" in r.text for r in bi_runs)


@pytest.mark.asyncio
async def test_inline_code(generator: WordGenerator) -> None:
    """`code` → run monospace."""
    result = await generator.generate(_params(content="Utiliser `print()` ici"))

    doc = Document(result.file_path)
    runs = doc.paragraphs[0].runs
    code_runs = [r for r in runs if r.font.name == "Courier New"]
    assert any("print()" in r.text for r in code_runs)


# === Lists ===


@pytest.mark.asyncio
async def test_bullet_list(generator: WordGenerator) -> None:
    """- item → List Bullet."""
    result = await generator.generate(_params(content="- Premier\n- Deuxieme"))

    doc = Document(result.file_path)
    bullets = [p for p in doc.paragraphs if "List Bullet" in p.style.name]
    assert len(bullets) == 2
    assert bullets[0].text == "Premier"
    assert bullets[1].text == "Deuxieme"


@pytest.mark.asyncio
async def test_numbered_list(generator: WordGenerator) -> None:
    """1. item → List Number."""
    result = await generator.generate(_params(content="1. Alpha\n2. Beta"))

    doc = Document(result.file_path)
    numbered = [p for p in doc.paragraphs if "List Number" in p.style.name]
    assert len(numbered) == 2
    assert numbered[0].text == "Alpha"


@pytest.mark.asyncio
async def test_nested_bullet_list(generator: WordGenerator) -> None:
    """Nested bullets → List Bullet 2, List Bullet 3."""
    content = "- Level 0\n  - Level 1\n    - Level 2"
    result = await generator.generate(_params(content=content))

    doc = Document(result.file_path)
    bullets = [p for p in doc.paragraphs if "List Bullet" in p.style.name]
    assert len(bullets) == 3
    assert bullets[0].style.name == "List Bullet"
    assert bullets[1].style.name == "List Bullet 2"
    assert bullets[2].style.name == "List Bullet 3"


@pytest.mark.asyncio
async def test_nested_numbered_list(generator: WordGenerator) -> None:
    """Nested numbered → List Number 2."""
    content = "1. Top\n  1. Sub"
    result = await generator.generate(_params(content=content))

    doc = Document(result.file_path)
    numbered = [p for p in doc.paragraphs if "List Number" in p.style.name]
    assert len(numbered) == 2
    assert numbered[0].style.name == "List Number"
    assert numbered[1].style.name == "List Number 2"


@pytest.mark.asyncio
async def test_mixed_nested_list(generator: WordGenerator) -> None:
    """Mixed bullet + numbered nesting."""
    content = "- Bullet\n  1. Numbered sub"
    result = await generator.generate(_params(content=content))

    doc = Document(result.file_path)
    assert any(p.style.name == "List Bullet" for p in doc.paragraphs)
    assert any(p.style.name == "List Number 2" for p in doc.paragraphs)


# === Tables ===


@pytest.mark.asyncio
async def test_table(generator: WordGenerator) -> None:
    """Markdown table → Word table."""
    content = "| Nom | Age |\n|-----|-----|\n| Alice | 30 |\n| Bob | 25 |"
    result = await generator.generate(_params(content=content))

    doc = Document(result.file_path)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # header + 2 data rows
    assert table.rows[0].cells[0].text == "Nom"
    assert table.rows[1].cells[0].text == "Alice"


# === Code blocks ===


@pytest.mark.asyncio
async def test_code_block(generator: WordGenerator) -> None:
    """```code``` → paragraphe monospace."""
    content = '```python\ndef hello():\n    print("hi")\n```'
    result = await generator.generate(_params(content=content))

    doc = Document(result.file_path)
    code_paras = [p for p in doc.paragraphs if p.runs and p.runs[0].font.name == "Courier New"]
    assert len(code_paras) >= 1


# === Blockquotes ===


@pytest.mark.asyncio
async def test_blockquote(generator: WordGenerator) -> None:
    """> text → Intense Quote style."""
    result = await generator.generate(_params(content="> Citation importante"))

    doc = Document(result.file_path)
    quotes = [p for p in doc.paragraphs if p.style.name == "Intense Quote"]
    assert len(quotes) == 1
    assert quotes[0].text == "Citation importante"


# === Page breaks ===


@pytest.mark.asyncio
async def test_horizontal_rule_page_break(generator: WordGenerator) -> None:
    """--- → page break."""
    result = await generator.generate(_params(content="Page 1\n\n---\n\nPage 2"))

    doc = Document(result.file_path)
    breaks = doc.element.findall(f".//{qn('w:br')}[@{qn('w:type')}='page']")
    assert len(breaks) >= 1


# === Metadata ===


@pytest.mark.asyncio
async def test_author_metadata(generator: WordGenerator) -> None:
    """author → propriété creator du document."""
    result = await generator.generate(_params(content="Test", author="John Doe"))

    doc = Document(result.file_path)
    assert doc.core_properties.author == "John Doe"


@pytest.mark.asyncio
async def test_title_metadata(generator: WordGenerator) -> None:
    """title → propriété title du document."""
    result = await generator.generate(_params(content="Test", title="Mon Rapport"))

    doc = Document(result.file_path)
    assert doc.core_properties.title == "Mon Rapport"


# === Result ===


@pytest.mark.asyncio
async def test_result_fields(generator: WordGenerator) -> None:
    """Le résultat contient tous les champs attendus."""
    result = await generator.generate(_params(filename="output.docx", content="Hello"))

    assert result.filename == "output.docx"
    assert result.file_path.endswith("output.docx")
    assert result.file_size_bytes > 0
    assert result.overwritten is False


@pytest.mark.asyncio
async def test_overwrite_flag(generator: WordGenerator) -> None:
    """Overwrite détecté quand le fichier existe déjà."""
    await generator.generate(_params(content="First"))
    result = await generator.generate(_params(content="Second"))
    assert result.overwritten is True


# === Errors ===


@pytest.mark.asyncio
async def test_filename_path_traversal(generator: WordGenerator) -> None:
    """Path traversal dans filename → WordGenerationError."""
    params = CreateWordParams.model_construct(
        filename="../evil.docx",
        content="test",
    )
    with pytest.raises(WordGenerationError):
        await generator.generate(params)


@pytest.mark.asyncio
async def test_output_too_large(generator: WordGenerator) -> None:
    """Fichier trop gros → WordGenerationError."""
    with patch.object(generator, "_max_output_size_mb", 0), pytest.raises(WordGenerationError):
        await generator.generate(_params(content="Test"))
